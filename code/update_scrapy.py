# -*- coding:utf-8 -*-

import datetime
import requests
import re
from lxml import etree
from docx import Document
import pymysql



def write(times):
    with open("update_time.txt", "a+") as f:
        time = datetime.datetime.now()
        time = datetime.datetime.strftime(time, "%Y-%m-%d %H:%M")
        f.write("The update time is %s," % str(times+1))
        f.write("and now time is %s.\n" % time)


def read():
    with open("update_time.txt", "r") as f:
        txt = f.read()
        pattern_date = "\d{4}-\d{2}-\d{2}"
        pattern_times = "is \d\,"
        date = (re.findall(pattern_date, txt))[-1]
        date = datetime.datetime.strptime(date, "%Y-%m-%d")

        times_01 =((re.findall(pattern_times, txt))[-1])
        times = int((re.findall("\d", times_01))[0])
    return date,times


def create_word(title,date,report,content,url, docx_title):
    file = Document()
    file.add_paragraph(title)
    file.add_paragraph(date)
    file.add_paragraph(report)
    file.add_paragraph(content)
    file.add_paragraph(url)

    file.save("E:\\Python\\web_scrapy_with_python\\china_us_trade\\document\\%s.docx" % docx_title)


def connent_mysql():
    db = pymysql.connect(host="localhost",
                         user='root',
                         port=3306,
                         db='chinaustrade',
                         password='******',
                         charset='utf8')
    return db


def HtmlDownloader():
    (date, times) = read()
    now = datetime.datetime.now()
    diff_update = (now - date).days
    pattern = "\d{6}/\d{2}"

    urls = []
    for i in range(1, 3):
        url = 'http://www.chinadaily.com.cn/world/china-us/page_%d.html' % i
        urls_list = HtmlDownloader_urls(url)
        for url in urls_list:
            a = re.search(pattern, url).group()
            a = datetime.datetime.strptime(a, "%Y%m/%d")
            diff_real = (now - a).days
            if diff_real <= diff_update :
                urls.append(url)
            else:
                print("已经存在于原来的数据库中")
        print("第%d页链接抓取完成" % i)
    write(times)
    return urls


def HtmlDownloader_urls(url):
    if url is None:
        return None
    user_agent = 'Mozilla/4.0 (compatible; MSIE 5.5; Windows NT)'
    headers = {'User-Agent': user_agent}
    r = requests.get(url, headers=headers)
    if r.status_code == 200:
        r.encoding = "utf-8"
    html = etree.HTML(r.text)
    urls_list = html.xpath('.//*[@class="mb10 tw3_01_2 "]/a/@href')
    return urls_list


def HtmlDownloader_text(url,i):
    print("正在抓取的网址为%s" % url)
    user_agent = 'Mozilla/4.0 (compatible; MSIE 5.5; Windows NT)'
    headers = {'User-Agent': user_agent}
    r = requests.get(url, headers=headers)
    if r.status_code == 200:
        r.encoding = "utf-8"
    html = etree.HTML(r.text)
    list_title = html.xpath(".//*[@id='lft-art']/h1")
    content_list = html.xpath(".//*[@id='Content']/p")
    report_list = html.xpath(".//*[@id='lft-art']/div[1]/span[1]")

    j = 1
    if len(list_title) == 0:
        a = str(j)
        title = a
        j += 1
    else:
        title = list_title[0].text

    content_part = ""
    for content_i in content_list:
        content01 = content_i.text
        if type(content01) == str:
            content_part = content_part + "\n" + content01
        else:
            print("正文非字符串类型")
    content = content_part

    if len(report_list) == 0:
        report = ""
    else:
        report = report_list[0].text

    pattern = "\d{6}/\d{2}"
    time = re.search(pattern, url).group()
    date = datetime.datetime.strptime(time,"%Y%m/%d")
    date = datetime.datetime.strftime(date, "%Y-%m-%d")
    pattern = "\d{4}\-\d{2}\-\d{2} \d{2}\:\d{2}"
    a = re.search(pattern, report)
    if a == None:
        convert_title = datetime.datetime.now()
    else:
        a = a.group()
        convert_title = datetime.datetime.strptime(a, "%Y-%m-%d %M:%S")
    docx_title = datetime.datetime.strftime(convert_title, "%Y-%m-%d-%M-%S")

    create_word(title, date, report, content, url, docx_title)
    print("正在写入第%d个word文档" % i)
    return [title, date, report, content, url]


def Html_text():
    i = 0

    urls = HtmlDownloader()
    db = connent_mysql()
    for url in urls:
        url = "http:" + url
        article = HtmlDownloader_text(url,i)

        print("正在执行第%d条记录插入语句" % i)
        sql = 'insert into chinadaily(title,date,report,content,url) values(%s,%s,%s,%s,%s)'
        db.cursor().execute(sql, article)
        i += 1

    try:
        #print("ok")
        db.commit()
    except:
        db.rollback()
    return "这真是愉快的一天"

if __name__ == "__main__":
    Html_text()
    print("更新结束,再见朋友!")


